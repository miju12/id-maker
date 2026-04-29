import streamlit as st
from PIL import Image, ImageDraw, ImageOps
import io
import math
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from pypdf import PdfReader, PdfWriter
from streamlit_cropper import st_cropper

st.set_page_config(page_title="Print & Document Pro", layout="centered", page_icon="🖨️")

st.title("🖨️ Print & Document Pro")

# Create tabs for the four tools
tab1, tab2, tab3, tab4 = st.tabs([
    "🪪 ID Single", 
    "🛂 Passport Maker", 
    "📇 ID Front & Back", 
    "🗜️ PDF/Photo Tools"
])

# ==========================================
# TAB 1: ID CARD FORMATTER (Single Side)
# ==========================================
with tab1:
    st.markdown("Scale and format your single-sided ID designs for high-resolution printing.")
    
    with st.expander("⚙️ ID Layout Settings", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            orientation = st.radio("Card Orientation", ["Horizontal (Landscape)", "Vertical (Portrait)"], key="radio_orient_1")
            apply_rounding = st.checkbox("Apply Rounded Corners", value=False, key="check_round_1")
            radius = st.slider("Corner Radius (px)", 0, 100, 30, disabled=not apply_rounding, key="slider_round_1")
        with col2:
            long_edge_mm = st.number_input("Long Edge (mm)", value=89.0, step=1.0, key="num_long_1")
            short_edge_mm = st.number_input("Short Edge (mm)", value=59.0, step=1.0, key="num_short_1")
            dpi = st.number_input("Print Resolution (DPI)", value=300, step=50, key="num_dpi_1")

    width_mm = long_edge_mm if "Horizontal" in orientation else short_edge_mm
    height_mm = short_edge_mm if "Horizontal" in orientation else long_edge_mm

    uploaded_id = st.file_uploader("Upload ID Card Design...", type=["jpg", "jpeg", "png"], key="id_uploader")

    if uploaded_id is not None:
        original_image = Image.open(uploaded_id).convert("RGBA")
        target_width_px = int((width_mm / 25.4) * dpi)
        target_height_px = int((height_mm / 25.4) * dpi)
        
        processed_image = original_image.resize((target_width_px, target_height_px), Image.Resampling.LANCZOS)
        
        if apply_rounding:
            mask = Image.new('L', processed_image.size, 0)
            draw = ImageDraw.Draw(mask)
            draw.rounded_rectangle([(0, 0), processed_image.size], radius=radius, fill=255)
            rounded_image = Image.new('RGBA', processed_image.size, (0, 0, 0, 0))
            rounded_image.paste(processed_image, (0, 0), mask=mask)
            processed_image = rounded_image

        st.image(processed_image, caption=f"Final Size: {width_mm} x {height_mm} mm @ {dpi} DPI", use_container_width=True)
        
        buf = io.BytesIO()
        processed_image.save(buf, format="PNG")
        st.download_button("Download Print-Ready ID Image", data=buf.getvalue(), file_name="print_ready_id.png", mime="image/png", type="primary", key="dl_1")


# ==========================================
# TAB 2: PASSPORT PHOTO MAKER (Updated for Multiple Files)
# ==========================================
with tab2:
    st.markdown("Auto-crop and densely pack multiple photos into an A4 Word document.")
    
    size_options = {
        "India Passport (35 x 45 mm)": (35, 45),
        "US Passport (51 x 51 mm)": (51, 51),
        "Stamp Size (20 x 25 mm)": (20, 25),
        "Custom Size": (None, None)
    }
    
    col1, col2 = st.columns(2)
    with col1:
        selected_size = st.selectbox("Select Passport Size", list(size_options.keys()), key="select_pass_2")
        quantity = st.number_input("Copies per Photo", min_value=1, max_value=200, value=7, step=1, key="num_qty_2")
    with col2:
        if selected_size == "Custom Size":
            pass_width_mm = st.number_input("Width (mm)", value=35.0, step=1.0, key="pass_w_2")
            pass_height_mm = st.number_input("Height (mm)", value=45.0, step=1.0, key="pass_h_2")
        else:
            pass_width_mm, pass_height_mm = size_options[selected_size]
            st.info(f"Dimensions: {pass_width_mm}mm x {pass_height_mm}mm")

    add_cut_lines = st.checkbox("Add Dotted Cutting Lines", value=True, key="check_cut_2")
    
    # Enabled multiple file uploads
    uploaded_passes = st.file_uploader("Upload Portrait Photos...", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="pass_uploader")

    if uploaded_passes:
        processed_images = []
        
        st.subheader("Processed Previews")
        # Display previews dynamically based on how many images were uploaded
        cols = st.columns(min(len(uploaded_passes), 4) if len(uploaded_passes) > 0 else 1)
        
        for idx, uploaded_pass in enumerate(uploaded_passes):
            pass_img = Image.open(uploaded_pass).convert("RGB")
            target_w_px = int((pass_width_mm / 25.4) * 300)
            target_h_px = int((pass_height_mm / 25.4) * 300)
            
            cropped_pass = ImageOps.fit(pass_img, (target_w_px, target_h_px), Image.Resampling.LANCZOS)
            
            if add_cut_lines:
                draw = ImageDraw.Draw(cropped_pass)
                dash_len = 15
                line_width = 3
                w, h = cropped_pass.size
                for x in range(0, w, dash_len * 2):
                    draw.line([(x, 0), (x + dash_len, 0)], fill="gray", width=line_width)
                    draw.line([(x, h-1), (x + dash_len, h-1)], fill="gray", width=line_width)
                for y in range(0, h, dash_len * 2):
                    draw.line([(0, y), (0, y + dash_len)], fill="gray", width=line_width)
                    draw.line([(w-1, y), (w-1, y + dash_len)], fill="gray", width=line_width)

            processed_images.append(cropped_pass)
            
            with cols[idx % len(cols)]:
                st.image(cropped_pass, use_container_width=True)

        total_photos = quantity * len(processed_images)
        cols_portrait = 200 // pass_width_mm
        rows_portrait = math.ceil(total_photos / cols_portrait) if cols_portrait > 0 else 999
        
        cols_landscape = 287 // pass_width_mm
        rows_landscape = math.ceil(total_photos / cols_landscape) if cols_landscape > 0 else 999

        best_orientation = "Landscape" if rows_landscape < rows_portrait else "Portrait"

        def generate_passport_doc(orientation):
            doc = Document()
            for section in doc.sections:
                if orientation == "Landscape":
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Mm(297)
                    section.page_height = Mm(210)
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Mm(210)
                    section.page_height = Mm(297)
                section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Mm(5)
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = 0 
            p.paragraph_format.space_after = Pt(2) 
            run = p.add_run()
            
            for img in processed_images:
                img_buffer = io.BytesIO()
                img.save(img_buffer, format="JPEG", quality=95)
                
                for _ in range(quantity):
                    img_buffer.seek(0) 
                    run.add_picture(img_buffer, width=Mm(pass_width_mm), height=Mm(pass_height_mm))
                    # Adds EXACTLY two spaces between photos
                    run.add_text("  ") 
                
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            return docx_buffer.getvalue()

        st.success(f"Calculated best layout: **{best_orientation} A4** to minimize paper loss for {total_photos} total photos.")
        st.download_button(
            "Generate Optimized Word Doc", 
            data=generate_passport_doc(best_orientation), 
            file_name="passports_print_ready.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            type="primary", 
            key="dl_2"
        )


# ==========================================
# TAB 3: ID FRONT & BACK COMBINER
# ==========================================
with tab3:
    st.markdown("Crop, merge, and perfectly align the front and back of an ID (Aadhaar, PAN, etc.) onto a single A4 page.")
    
    with st.expander("⚙️ Layout Settings", expanded=True):
        st.caption("Perform a free form crop on the images below. We will automatically stretch the result to match these fixed printable dimensions.")
        col_set1, col_set2 = st.columns(2)
        with col_set1:
            combine_width = st.number_input("Target Width (mm)", value=86.0, step=1.0, key="combine_w_3")
        with col_set2:
            combine_height = st.number_input("Target Height (mm)", value=54.0, step=1.0, key="combine_h_3")
    
    st.divider()

    col_f, col_b = st.columns(2)
    
    front_cropped = None
    back_cropped = None

    def get_high_res_crop(raw_image, key, box_color):
        w, h = raw_image.size
        disp_w = 350 
        
        if w > disp_w:
            disp_h = int(h * (disp_w / w))
            disp_img = raw_image.resize((disp_w, disp_h), Image.Resampling.LANCZOS)
            scale = w / disp_w
        else:
            disp_img = raw_image
            scale = 1.0
            
        st.caption("Drag the box for free form crop")
        box = st_cropper(disp_img, realtime_update=True, box_color=box_color, aspect_ratio=None, return_type='box', key=key)
        
        if box:
            left = int(box['left'] * scale)
            top = int(box['top'] * scale)
            right = int((box['left'] + box['width']) * scale)
            bottom = int((box['top'] + box['height']) * scale)
            return raw_image.crop((left, top, right, bottom))
        return None

    with col_f:
        st.subheader("Front Side")
        front_file = st.file_uploader("Upload Front Side", type=["jpg", "jpeg", "png"], key="f_up_3")
        if front_file:
            front_img_raw = Image.open(front_file).convert("RGB")
            front_cropped = get_high_res_crop(front_img_raw, key="crop_f_3", box_color='#FF0000')

    with col_b:
        st.subheader("Back Side")
        back_file = st.file_uploader("Upload Back Side", type=["jpg", "jpeg", "png"], key="b_up_3")
        if back_file:
            back_img_raw = Image.open(back_file).convert("RGB")
            back_cropped = get_high_res_crop(back_img_raw, key="crop_b_3", box_color='#0088FF')

    if front_cropped and back_cropped:
        st.divider()
        st.subheader("Processed Printable Preview (Auto-Resized to Fixed Dimensions)")
        col_prev1, col_prev2 = st.columns(2)
        
        target_w_px_3 = int((combine_width / 25.4) * 300)
        target_h_px_3 = int((combine_height / 25.4) * 300)
        
        front_final_3 = front_cropped.resize((target_w_px_3, target_h_px_3), Image.Resampling.LANCZOS)
        back_final_3 = back_cropped.resize((target_w_px_3, target_h_px_3), Image.Resampling.LANCZOS)
            
        with col_prev1:
            st.image(front_final_3, caption=f"Front (Scaled to {combine_width}x{combine_height}mm)", use_container_width=True)
        with col_prev2:
            st.image(back_final_3, caption=f"Back (Scaled to {combine_width}x{combine_height}mm)", use_container_width=True)
            
        def generate_combined_doc_3():
            doc = Document()
            for section in doc.sections:
                section.page_width = Mm(210)
                section.page_height = Mm(297)
                section.left_margin = Mm(15)
                section.top_margin = Mm(15)
                
            f_buf_3 = io.BytesIO()
            b_buf_3 = io.BytesIO()
            front_final_3.save(f_buf_3, format="JPEG", quality=95)
            back_final_3.save(b_buf_3, format="JPEG", quality=95)
            
            p1_3 = doc.add_paragraph()
            r1_3 = p1_3.add_run()
            f_buf_3.seek(0)
            r1_3.add_picture(f_buf_3, width=Mm(combine_width), height=Mm(combine_height))
            
            p_space_3 = doc.add_paragraph()
            p_space_3.paragraph_format.space_after = Pt(20)
            
            p2_3 = doc.add_paragraph()
            r2_3 = p2_3.add_run()
            b_buf_3.seek(0)
            r2_3.add_picture(b_buf_3, width=Mm(combine_width), height=Mm(combine_height))
            
            doc_out_3 = io.BytesIO()
            doc.save(doc_out_3)
            return doc_out_3.getvalue()
            
        st.download_button(
            label="Download A4 Printable Document",
            data=generate_combined_doc_3(),
            file_name="ID_Front_and_Back.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
            key="dl_3"
        )

# ==========================================
# TAB 4: PHOTO & PDF TOOLS
# ==========================================
with tab4:
    st.markdown("Compress photos, compress PDFs, or merge multiple images into a single PDF document.")
    
    tool_choice = st.radio("Select Tool", ["Compress Photo", "Merge Images to PDF", "Compress PDF"], key="radio_tool_4")
    
    if tool_choice == "Compress Photo":
        comp_img = st.file_uploader("Upload Image to Compress", type=["jpg", "jpeg", "png"], key="comp_up_4")
        if comp_img:
            quality = st.slider("JPEG Quality (Lower = Smaller File)", 10, 100, 60, key="slider_qual_4")
            img = Image.open(comp_img).convert("RGB")
            
            out_buf = io.BytesIO()
            img.save(out_buf, format="JPEG", optimize=True, quality=quality)
            size_kb = len(out_buf.getvalue()) / 1024
            
            st.success(f"Compressed Size: {size_kb:.2f} KB")
            st.download_button("Download Compressed Image", data=out_buf.getvalue(), file_name="compressed_image.jpg", mime="image/jpeg", type="primary", key="dl_photo_4")

    elif tool_choice == "Merge Images to PDF":
        merge_imgs = st.file_uploader("Upload Multiple Images", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="merge_up_4")
        if merge_imgs:
            st.info(f"{len(merge_imgs)} images loaded.")
            if st.button("Merge into PDF", type="primary", key="btn_merge_4"):
                img_list = []
                for upload in merge_imgs:
                    img = Image.open(upload).convert("RGB")
                    img_list.append(img)
                
                if img_list:
                    pdf_buf = io.BytesIO()
                    img_list[0].save(pdf_buf, format="PDF", save_all=True, append_images=img_list[1:])
                    st.download_button("Download Merged PDF", data=pdf_buf.getvalue(), file_name="merged_photos.pdf", mime="application/pdf", type="primary", key="dl_pdf_merge_4")

    elif tool_choice == "Compress PDF":
        comp_pdf = st.file_uploader("Upload PDF to Compress", type=["pdf"], key="comp_pdf_4")
        if comp_pdf:
            if st.button("Compress PDF", type="primary", key="btn_pdf_4"):
                reader = PdfReader(comp_pdf)
                writer = PdfWriter()

                for page in reader.pages:
                    writer.add_page(page)
                
                for page in writer.pages:
                    page.compress_content_streams()
                
                pdf_out = io.BytesIO()
                writer.write(pdf_out)
                size_kb = len(pdf_out.getvalue()) / 1024
                
                st.success(f"Compressed PDF Size: {size_kb:.2f} KB")
                st.download_button("Download Compressed PDF", data=pdf_out.getvalue(), file_name="compressed_document.pdf", mime="application/pdf", type="primary", key="dl_pdf_comp_4")
